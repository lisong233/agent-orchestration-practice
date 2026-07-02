"""
文档 I/O 层 — 统一的文档→文本转换入口。

解析策略（优先级降级）：
  1. doc-read CLI（本地开发，OfficeCLI HTML→Markdown，最准）
  2. LibreOffice headless（Docker，docx→text，高保真）
  3. python-docx（兜底，读段落+表格纯文本）
"""
import subprocess
import os
import tempfile
from pathlib import Path


def _via_docread(path: Path) -> str | None:
    """方式 A：doc-read CLI（本地开发环境）"""
    try:
        result = subprocess.run(
            ["doc-read", str(path)],
            capture_output=True, text=True, encoding="utf-8", timeout=30,
        )
        if result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _via_libreoffice(path: Path) -> str | None:
    """方式 B：LibreOffice headless（Docker 环境，Linux）"""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(path)],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0:
                # LibreOffice 输出文件名为 原文件名.txt
                out_name = path.stem + ".txt"
                out_path = os.path.join(tmpdir, out_name)
                if os.path.exists(out_path):
                    with open(out_path, encoding="utf-8", errors="replace") as f:
                        content = f.read()
                    if content.strip():
                        return content
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _via_python_docx(path: Path) -> str | None:
    """方式 C：python-docx（纯 Python 兜底）"""
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        # 段落
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 表格：保留结构
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    text = cell.text.strip().replace('\n', ' ')
                    cells.append(text)
                parts.append(" | ".join(cells))

        return "\n".join(parts)
    except ImportError:
        pass
    return None


def to_text(file_path: str | Path) -> str:
    """将文档转换为纯文本。支持 .txt / .docx / .doc / .pdf。"""
    path = Path(file_path)

    # .txt 直接读
    if path.suffix.lower() == ".txt":
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()

    # 二进制格式 → 逐级降级
    if path.suffix.lower() in (".docx", ".doc", ".pdf"):
        # A：doc-read CLI
        text = _via_docread(path)
        if text:
            return text

        # B：LibreOffice headless
        text = _via_libreoffice(path)
        if text:
            return text

        # C：python-docx
        text = _via_python_docx(path)
        if text:
            return text

        # D：终极兜底——当纯文本读
        try:
            with open(path, encoding="utf-8", errors="replace") as f:
                content = f.read()
                if len(content.strip()) > 50:
                    return content
        except Exception:
            pass

        raise RuntimeError(
            f"无法解析 {path.suffix} 文件。"
            f"\n  开发环境：pip install doc-read"
            f"\n  Docker 环境：需 LibreOffice 或 python-docx"
        )

    raise ValueError(f"不支持的文件格式: {path.suffix}")


def detect_doc_type_from_path(file_path: str | Path) -> str | None:
    """从文件扩展名推断文档类型"""
    suffix = Path(file_path).suffix.lower()
    if suffix in (".docx", ".doc", ".pdf", ".txt"):
        return None
    return None
